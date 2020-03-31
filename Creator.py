import re
import datetime
import os
import comtypes.client
from docx import Document

stawkaPL = {
    '12' : 'dwanaście',
    '13' : 'trzynaście',
    '14' : 'czternaście',
    '15' : 'piętnaście',
    '16' : 'szesnaście',
    '17' : 'siedemnaście',
    '18' : 'osiemnaście',
    '19' : 'dziewiętnaście',
    '20' : 'dwadzieścia',
}

stawkaANG = {
    '12' : 'twelve',
    '13' : 'thirteen',
    '14' : 'fourteen',
    '15' : 'fifteen',
    '16' : 'sixteen',
    '17' : 'seventeen',
    '18' : 'eightteen',
    '19' : 'nineteen',
    '20' : 'twenty',
}

now = datetime.datetime.now()
day = str(now.day)
month = str(now.month)
year = str(now.year)
data = day + '.' + month + '.' + year

def rate_in_word(stawka, replace6):
    wybor = replace6
    while True:
        try:
            stawka[wybor]
            break
        except KeyError:
            print('ERROR! Podano stawke z złego zakresu!')
            exit()
            break
    return stawka[wybor]


def replace_word(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_word(cell, regex, replace)

def word_to_pdf(_in, _out):
    pdf_format_key = 17
    file_in = os.path.abspath(_in)
    file_out = os.path.abspath(_out)
    worddoc = comtypes.client.CreateObject('Word.Application')
    doc = worddoc.Documents.Open(file_in)
    doc.SaveAs(file_out, FileFormat = pdf_format_key)
    doc.Close()
    worddoc.Quit()

regex1 = re.compile(r"Dzis") #Przypisanie do zmiennej, szukanego slowa "Dzis"
regex2 = re.compile(r"Nazwisko") #Przypisanie do zmiennej, szukanego slowa "Nazwisko"
regex3 = re.compile(r"1234567890") #Przypisanie do zmiennej, szukanego slowa "1234567890"
regex4 = re.compile(r"DWADZIESCIALITERWERS") #Przypisanie do zmiennej, szukanego slowa "DWADZIESCIALITERWERS"
regex5 = re.compile(r"ODDO") #Przypisanie do zmiennej, szukanego slowa "ODDO"
regex6 = re.compile(r"95") #Przypisanie do zmiennej, szukanego slowa "95"
regex7 = re.compile(r"PLslownie") #Przypisanie do zmiennej, szukanego slowa "PLslownie"
regex8 = re.compile(r"ANGslownie") #Przypisanie do zmiennej, szukanego slowa "ANGslownie"

replace1 = data #Wpisz date
replace2 = input(r"Wpisz imie i nazwisko: ") #Wpisz imie i nazwisko
replace3 = input(r"Nr paszportu: ")
replace4 = input(r"Statek: ")
replace5 = input(r"Czas trwania projektu: ")
replace6 = input(r"Stawka: ")
replace7 = rate_in_word(stawkaPL, replace6)
replace8 = rate_in_word(stawkaANG, replace6)

filename = "Umowa.docx" #Nazwa szablnu

doc1 = Document(filename)
replace_word(doc1, regex1, replace1)
replace_word(doc1, regex2, replace2)
replace_word(doc1, regex3, replace3)
replace_word(doc1, regex4, replace4)
replace_word(doc1, regex5, replace5)
replace_word(doc1, regex6, replace6)
replace_word(doc1, regex7, replace7)
replace_word(doc1, regex8, replace8)

name_pdf_file = replace2 + '_' + replace4  + '.pdf'

doc1.save('wynik.docx') #Docelowy plik gotowy
word_to_pdf('wynik.docx', name_pdf_file)



